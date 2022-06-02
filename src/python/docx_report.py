import os

import numpy as np

from docx import Document
NUM_FOR_ROUND = 5


def create_document(heading: str):
    """
    Создает документ в формате docx
    :param heading: заголовок документа
    :return: документ в формате docx
    """
    document = Document()
    document.add_heading(heading, 1)
    return document


def save_document(document: Document, path: str, name: str):
    """
    Сохраняет документ в формате docx
    :param document: документ
    :param path: путь к папке, в которую будут сохранен документ
    :param name: название документа
    :return: None
    """
    report_path = os.path.join(path, name)
    document.save(report_path)


def num_to_str(num) -> str:
    """
    Приводит строку к числу
    :param num: число
    :return: строка
    """
    if isinstance(num, str):
        return num
    return str(np.round(num, NUM_FOR_ROUND))


# Заполняет таблицу, которая будет сохранения в DOCX-файле
def fill_table_for_report(document, data):
    """Создает таблицу в документе и заполняет ее."""
    table = document.add_table(rows=len(data), cols=len(data[0]), style='Table Grid')
    for i, line in enumerate(data):
        for j, num in enumerate(line):
            table.rows[i].cells[j].text = num_to_str(num)
    return table


# Добавляет таблицу в DOCX-файл
def add_table_to_report(heading: str, data: list[list], report: Document):
    report.add_heading(heading, 2)
    fill_table_for_report(report, data)
